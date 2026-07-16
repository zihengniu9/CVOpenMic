"""解析器安全边界与格式兼容性测试。"""

import io
import unittest
from unittest.mock import patch

from docx import Document
from PIL import Image
from PyPDF2 import PdfWriter

import parser
import config


def _pdf_bytes(page_count: int) -> bytes:
    buffer = io.BytesIO()
    writer = PdfWriter()
    for _ in range(page_count):
        writer.add_blank_page(width=595, height=842)
    writer.write(buffer)
    return buffer.getvalue()


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("第一段：产品经理，负责核心业务规划与跨团队协作。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目成果"
    table.cell(0, 1).text = "推动转化率提升并完成关键版本交付"
    document.add_paragraph("最后一段：熟悉数据分析、用户研究和项目管理。")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class ParserTests(unittest.TestCase):
    def test_resume_text_limit_comes_from_config(self):
        self.assertEqual(parser.MAX_TEXT_CHARS, config.MAX_RESUME_CHARS)

    def test_legacy_doc_is_explicitly_rejected(self):
        with self.assertRaisesRegex(ValueError, r"不支持旧版 \.doc"):
            parser.detect_format("resume.doc")

    def test_invalid_pdf_signature_is_rejected_before_parsing(self):
        with self.assertRaisesRegex(ValueError, "不是有效的 PDF"):
            parser.parse_resume(b"this is not a pdf", "resume.pdf")

    def test_blank_pdf_reports_scanned_document(self):
        with self.assertRaisesRegex(ValueError, "扫描件"):
            parser.parse_pdf(_pdf_bytes(1))

    def test_pdf_page_limit_comes_from_config(self):
        with patch.object(parser, "MAX_PAGES", 1):
            with self.assertRaisesRegex(ValueError, "页数超过限制"):
                parser.parse_pdf(_pdf_bytes(2))

    def test_docx_blocks_keep_document_order(self):
        text, fmt = parser.parse_resume(_docx_bytes(), "resume.docx")

        self.assertEqual(fmt, "docx")
        self.assertLess(text.index("第一段"), text.index("项目成果"))
        self.assertLess(text.index("项目成果"), text.index("最后一段"))

    def test_invalid_zip_is_not_accepted_as_docx(self):
        fake_zip = b"PK\x03\x04not-a-real-docx"
        with self.assertRaisesRegex(ValueError, "不是有效的 DOCX"):
            parser.parse_resume(fake_zip, "resume.docx")

    def test_gb18030_text_is_supported(self):
        content = (
            "张三，产品经理。负责用户研究、需求分析和版本规划，"
            "推动多个跨团队项目按期上线，并持续优化核心业务流程。"
        )
        text, fmt = parser.parse_resume(content.encode("gb18030"), "resume.txt")

        self.assertEqual(fmt, "txt")
        self.assertIn("产品经理", text)

    def test_utf8_bom_text_is_supported(self):
        content = (
            "李四，数据分析师。熟悉 SQL、Python 和指标体系建设，"
            "能够独立完成业务分析、实验复盘和可视化报告交付。"
        )
        text, fmt = parser.parse_resume(content.encode("utf-8-sig"), "resume.txt")

        self.assertEqual(fmt, "txt")
        self.assertTrue(text.startswith("李四"))

    def test_file_size_limit_is_enforced(self):
        with patch.object(parser, "MAX_FILE_SIZE_MB", 0):
            with self.assertRaisesRegex(ValueError, "文件太大"):
                parser.parse_resume(b"not empty", "resume.txt")

    def test_text_character_limit_is_enforced(self):
        with patch.object(parser, "MAX_TEXT_CHARS", 20):
            with self.assertRaisesRegex(ValueError, "文本过长"):
                parser.parse_resume(("有效经历" * 20).encode("utf-8"), "resume.txt")

    def test_image_pixel_limit_is_checked_before_ocr(self):
        buffer = io.BytesIO()
        Image.new("RGB", (11, 11), "white").save(buffer, format="PNG")

        with patch.object(parser, "MAX_IMAGE_PIXELS", 100):
            with self.assertRaisesRegex(ValueError, "像素过大"):
                parser.parse_resume(buffer.getvalue(), "resume.png")

    def test_image_extension_must_match_signature(self):
        buffer = io.BytesIO()
        Image.new("RGB", (8, 8), "white").save(buffer, format="PNG")

        with self.assertRaisesRegex(ValueError, "扩展名与实际格式不一致"):
            parser.parse_resume(buffer.getvalue(), "resume.jpg")


if __name__ == "__main__":
    unittest.main()
