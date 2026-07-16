"""Export CVOpenMic results to recruiter-friendly files."""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


_LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def _clean_inline_markdown(text: str) -> str:
    """Remove the small Markdown subset commonly produced by the model."""
    text = _LINK_RE.sub(r"\1", text)
    text = re.sub(r"(\*\*|__|`)", "", text)
    return text.strip()


def create_docx_bytes(markdown_text: str, title: str = "优化简历") -> bytes:
    """Create a simple ATS-friendly DOCX and return its bytes."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.core_properties.title = title

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line or line in {"---", "***", "___"}:
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            document.add_heading(_clean_inline_markdown(heading.group(2)), level=level)
            continue

        bullet = re.match(r"^[-*+]\s+(.+)$", line)
        if bullet:
            document.add_paragraph(
                _clean_inline_markdown(bullet.group(1)), style="List Bullet"
            )
            continue

        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if numbered:
            document.add_paragraph(
                _clean_inline_markdown(numbered.group(1)), style="List Number"
            )
            continue

        document.add_paragraph(_clean_inline_markdown(line.lstrip("> ")))

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
